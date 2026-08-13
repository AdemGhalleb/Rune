"""DOCX paragraph block extractor implementation using python-docx."""

import logging
from pathlib import Path

import docx

from app.ai.extraction.base import (
    BaseExtractor,
    ExtractionResult,
    SegmentData,
    compute_text_hash,
)

logger = logging.getLogger(__name__)


class DocxExtractor(BaseExtractor):
    """DOCX extractor producing paragraph-block segments."""

    def extract(self, file_path: Path) -> ExtractionResult:
        doc = docx.Document(str(file_path))
        segments: list[SegmentData] = []
        all_text_parts: list[str] = []

        # Extract non-empty paragraphs and tables as paragraph blocks
        blocks: list[str] = []

        for paragraph in doc.paragraphs:
            txt = paragraph.text.strip()
            if txt:
                blocks.append(txt)

        for table in doc.tables:
            table_lines: list[str] = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_lines.append(row_text)
            if table_lines:
                blocks.append("\n".join(table_lines))

        # If no paragraphs or tables yielded text, provide empty or single block
        if not blocks:
            blocks = [""]

        for idx, block_text in enumerate(blocks):
            segments.append(
                SegmentData(
                    segment_index=idx,
                    segment_type="docx_paragraph_block",
                    text=block_text,
                    char_count=len(block_text),
                    page_number=None,
                )
            )
            all_text_parts.append(block_text)

        logger.info(
            "Extracted DOCX %s (%d blocks, %d chars)",
            file_path.name,
            len(segments),
            sum(s.char_count for s in segments),
        )

        combined_text = "\n\n".join(all_text_parts)
        text_hash = compute_text_hash(combined_text)

        return ExtractionResult(
            segments=segments,
            extracted_text_hash=text_hash,
            has_partial_errors=False,
        )
